import os
import sys
import subprocess
import ctypes
import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from tkcalendar import Calendar
from docx import Document
import win32com.client

# 要处理的目标文件名（严格匹配）
target_filenames = [
    "3、网络安全等级保护前期调研表.docx",
    "5-2、附件：测评方案评审意见及反馈记录.docx",
    "7、风险告知书.docx",
    "9、现场测评授权书（盖章）.docx",
    "12、文档交接单.docx",
    "13、测评现场记录表.docx",
    "16-1、网络安全等级保护测评报告评审表.docx",
    "16-2、附件：测评报告评审意见及反馈记录.docx",
    "19、客户满意度调查表(需盖章).docx",
    "20、电子版测评报告申请承诺书（第三方）.docx",
    "20、电子版测评报告申请承诺书（业主盖章）.docx",
    "1、项目归档清单.docx",
    "8、现场（末次）会议签到表.docx",
    "8、现场（首次）会议签到表.docx",
    "10、漏洞扫描和渗透测试授权书（需盖章）.docx",
    "11、漏洞扫描和渗透测试确认书.docx",
    "14、离场确认书.docx",
    "15、设备软件使用情况表.docx",
    "5、测评方案评审表.docx",
    "10、自愿放弃漏洞扫描及渗透测试声明（需盖章）.docx",
    "2、项目计划书（工作方案）二级非信创+渗透.docx",
    "2、项目计划书（工作方案）二级非信创+无渗透.docx",
    "2、项目计划书（工作方案）二级信创+渗透.docx",
    "2、项目计划书（工作方案）二级信创+无渗透.docx",
    "2、项目计划书（工作方案）三级.docx",
    "2、项目计划书（工作方案） 三级+上海.docx",
    "17、测评报告签收确认书.docx",
    "17、测评报告签收确认书（上海）.docx",
    "1、项目归档清单（仅扫描）.docx",
    "8、现场（末次）会议签到表（仅放弃渗透）.docx",
    "8、现场（首次）会议签到表（仅放弃渗透）.docx",
    "10、漏洞扫描和验证测试授权书（需盖章）.docx",
    "11、漏洞扫描和验证测试确认书.docx",
    "14、离场确认书（仅漏扫）.docx",
    "15、设备软件使用情况表（仅扫描）.docx",
    "1、项目归档清单（放弃漏扫渗透）.docx",
    "5、测评方案评审表+放弃漏扫渗透.docx",
    "8、现场（末次）会议签到表（放弃漏扫渗透）.docx",
    "8、现场（首次）会议签到表（放弃漏扫渗透）.docx",
    "14、离场确认书（放弃扫描+渗透）.docx",
    "15、设备软件使用情况表（放弃漏扫渗透）.docx",
]

# 替换项目配置
replace_items = [
    {"label": "合同号", "key": "AG202501-XXX", "type": "text"},
    {"label": "单位名称", "key": "XX单位", "type": "text"},
    {"label": "系统名称", "key": "XX系统", "type": "text"},
    {"label": "联系人", "key": "LXR", "type": "text"},
    {"label": "联系方式", "key": "LXFS", "type": "text"},
    {"label": "系统等级", "key": "S2A2G2", "type": "select", "options": ["", "S2A2G2", "S3A3G3", "S3A2G3", "S2A3G3"]},
    {"label": "进场时间", "key": "JCSJ", "type": "date"},
    {"label": "离场时间", "key": "LCSJ", "type": "date"},
    {"label": "备案编号", "key": "BABH", "type": "text"},
]

penetration_officers = ["胡乐华", "王辉", "戚建军"]

def replace_text_in_runs_preserve_style(runs, old_text, new_text):
    for run in runs:
        if old_text in run.text:
            run.text = run.text.replace(old_text, new_text)

def replace_in_docx(doc, replace_dict):
    for para in doc.paragraphs:
        for old, new in replace_dict.items():
            replace_text_in_runs_preserve_style(para.runs, old, new)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for old, new in replace_dict.items():
                        replace_text_in_runs_preserve_style(para.runs, old, new)

def select_folder():
    folder_selected = filedialog.askdirectory(title="请选择输入文件夹")
    if folder_selected:
        folder_path.set(folder_selected)

def select_output_folder():
    folder_selected = filedialog.askdirectory(title="请选择输出文件夹")
    if folder_selected:
        output_folder_path.set(folder_selected)

def open_calendar_popup(var, entry_widget):
    popup = tk.Toplevel()
    popup.title("请选择日期")
    popup.grab_set()
    cal = Calendar(popup, locale='zh_CN')
    cal.pack(padx=10, pady=10)

    def on_ok():
        selected_date = cal.selection_get()
        formatted_date = f"{selected_date.year}年{selected_date.month}月{selected_date.day}日"
        var.set(formatted_date)
        entry_widget.delete(0, tk.END)
        entry_widget.insert(0, formatted_date)
        popup.destroy()

    btn_ok = tk.Button(popup, text="确定", command=on_ok)
    btn_ok.pack(pady=5)
    popup.mainloop()

def docx_to_pdf(input_path, output_path):
    try:
        try:
            wps = win32com.client.Dispatch('Kwps.Application')
            wps.Visible = False
            doc = wps.Documents.Open(input_path)
            doc.ExportAsFixedFormat(output_path, 17)
            doc.Close()
            wps.Quit()
            return "WPS"
        except Exception:
            pass

        try:
            word = win32com.client.Dispatch('Word.Application')
            word.Visible = False
            doc = word.Documents.Open(input_path)
            doc.SaveAs(output_path, FileFormat=17)
            doc.Close()
            word.Quit()
            return "Office"
        except Exception:
            pass

        messagebox.showwarning(
            "未生成PDF",
            "未检测到 WPS 或 Microsoft Office，无法生成PDF文件"
        )
        return "无"
    except Exception as e:
        messagebox.showerror("错误", f"生成PDF时出错：\n{e}")
        return "错误"

def batch_convert_pdf(output_folder, files_to_convert):
    for f in files_to_convert:
        if f.lower().endswith(".docx"):
            input_path = os.path.join(output_folder, f)
            pdf_path = os.path.splitext(input_path)[0] + ".pdf"
            try:
                used_app = docx_to_pdf(input_path, pdf_path)
                print(f"使用了 {used_app} 将 {f} 转换成 {os.path.basename(pdf_path)}")
            except Exception as e:
                print(f"转换文件 {f} 失败：{e}")

def open_folder(path):
    if os.name == 'nt':
        os.startfile(path)
    elif sys.platform == "darwin":
        subprocess.Popen(["open", path])
    else:
        subprocess.Popen(["xdg-open", path])

def on_close():
    root.destroy()
    try:
        hwnd = ctypes.windll.kernel32.GetConsoleWindow()
        if hwnd != 0:
            ctypes.windll.user32.PostMessageW(hwnd, 0x0010, 0, 0)  # WM_CLOSE
    except:
        pass

def update_penetration_officer_row(*args):
    val = penetration_var.get()
    if val == "渗透+扫描":
        penetration_officer_label.grid()
        penetration_officer_cb.grid()
    else:
        penetration_officer_label.grid_remove()
        penetration_officer_cb.grid_remove()
        penetration_officer_var.set("")

def determine_files(region_val, system_level_val, is_xinchuan_val, penetration_val):
    # 默认基础文件（严格匹配）
    base_files = [
        "3、网络安全等级保护前期调研表.docx",
        "5-2、附件：测评方案评审意见及反馈记录.docx",
        "7、风险告知书.docx",
        "9、现场测评授权书（盖章）.docx",
        "12、文档交接单.docx",
        "13、测评现场记录表.docx",
        "16-1、网络安全等级保护测评报告评审表.docx",
        "16-2、附件：测评报告评审意见及反馈记录.docx",
        "19、客户满意度调查表(需盖章).docx",
        "20、电子版测评报告申请承诺书（第三方）.docx",
        "20、电子版测评报告申请承诺书（业主盖章）.docx",
    ]
    files = base_files.copy()

    if penetration_val == "渗透+扫描":
        files += [
            "1、项目归档清单.docx",
            "8、现场（末次）会议签到表.docx",
            "8、现场（首次）会议签到表.docx",
            "10、漏洞扫描和渗透测试授权书（需盖章）.docx",
            "11、漏洞扫描和渗透测试确认书.docx",
            "14、离场确认书.docx",
            "15、设备软件使用情况表.docx",
        ]
    elif penetration_val == "仅扫描":
        files += [
            "1、项目归档清单（仅扫描）.docx",
            "5、测评方案评审表.docx",
            "8、现场（末次）会议签到表（仅放弃渗透）.docx",
            "8、现场（首次）会议签到表（仅放弃渗透）.docx",
            "10、漏洞扫描和验证测试授权书（需盖章）.docx",
            "11、漏洞扫描和验证测试确认书.docx",
            "14、离场确认书（仅漏扫）.docx",
            "15、设备软件使用情况表（仅扫描）.docx",
        ]
    elif penetration_val == "放弃漏扫渗透":
        files += [
            "1、项目归档清单（放弃漏扫渗透）.docx",
            "5、测评方案评审表+放弃漏扫渗透.docx",
            "8、现场（末次）会议签到表（放弃漏扫渗透）.docx",
            "8、现场（首次）会议签到表（放弃漏扫渗透）.docx",
            "10、自愿放弃漏洞扫描及渗透测试声明（需盖章）.docx",
            "14、离场确认书（放弃扫描+渗透）.docx",
            "15、设备软件使用情况表（放弃漏扫渗透）.docx",
        ]

    if region_val == "浙江":
        files.append("17、测评报告签收确认书.docx")
    elif region_val == "上海":
        files.append("17、测评报告签收确认书（上海）.docx")

    if system_level_val == "S2A2G2":
        if is_xinchuan_val == "非信创":
            if penetration_val == "渗透+扫描":
                files.append("2、项目计划书（工作方案）二级非信创+渗透.docx")
            else:
                files.append("2、项目计划书（工作方案）二级非信创+无渗透.docx")
        elif is_xinchuan_val == "信创":
            if penetration_val == "渗透+扫描":
                files.append("2、项目计划书（工作方案）二级信创+渗透.docx")
            else:
                files.append("2、项目计划书（工作方案）二级信创+无渗透.docx")
    else:
        if region_val == "浙江":
            files.append("2、项目计划书（工作方案）三级.docx")
        elif region_val == "上海":
            files.append("2、项目计划书（工作方案） 三级+上海.docx")

    # 去重
    files = list(dict.fromkeys(files))
    return files

def start_replace():
    folder = folder_path.get()
    output_folder = output_folder_path.get()

    # 验证必选项
    if not region_var.get():
        messagebox.showerror("错误", "请选择地区")
        return
    if not xinchuan_var.get():
        messagebox.showerror("错误", "请选择是否信创")
        return
    if not penetration_var.get():
        messagebox.showerror("错误", "请选择是否渗透")
        return

    if not folder or not os.path.exists(folder):
        messagebox.showerror("错误", "请选择有效的输入文件夹路径")
        return
    if not output_folder or not os.path.exists(output_folder):
        messagebox.showerror("错误", "请选择有效的输出文件夹路径")
        return

    replace_dict = {}
    for i, item in enumerate(replace_items):
        if check_vars[i].get():
            val = entry_vars[i].get().strip()
            if val:
                replace_dict[item["key"]] = val

    if not replace_dict:
        messagebox.showwarning("提示", "没有启用任何替换项或未输入替换内容")
        return

    # 获取系统等级选择值
    system_level_val = entry_vars[5].get().strip()  # 替换项里系统等级索引5

    # 计算需要处理的文件列表
    files_to_process = determine_files(region_var.get(), system_level_val, xinchuan_var.get(), penetration_var.get())

    count = 0
    for filename in os.listdir(folder):
        if not filename.lower().endswith(".docx"):
            continue
        if filename not in files_to_process:
            continue

        full_input_path = os.path.join(folder, filename)
        try:
            doc = Document(full_input_path)
        except Exception as e:
            messagebox.showerror("错误", f"打开文件失败: {filename}\n{e}")
            continue

        replace_in_docx(doc, replace_dict)

        full_output_path = os.path.join(output_folder, filename)
        doc.save(full_output_path)
        count += 1

    messagebox.showinfo("完成", f"替换完成，共处理 {count} 个文件，开始批量转换PDF...")

    batch_convert_pdf(output_folder, files_to_process)

    # 自动打开输出文件夹
    open_folder(output_folder)

    # 结束对话框
    dlg = tk.Toplevel()
    dlg.title("完成")
    tk.Label(dlg, text=f"PDF转换完成，文件输出至：\n{output_folder}").pack(padx=20, pady=10)
    btn_open = tk.Button(dlg, text="打开文件夹", command=lambda: open_folder(output_folder), width=15)
    btn_open.pack(side="left", padx=20, pady=10)
    btn_close = tk.Button(dlg, text="关闭程序", command=on_close, width=15)
    btn_close.pack(side="right", padx=20, pady=10)
    dlg.grab_set()
    dlg.mainloop()

def check_all():
    for var in check_vars:
        var.set(True)

def uncheck_all():
    for var in check_vars:
        var.set(False)

root = tk.Tk()
root.title("Word批量替换工具（多文件+自动转PDF）")

folder_path = tk.StringVar()
output_folder_path = tk.StringVar()

# 地区 必选，默认空
region_var = tk.StringVar(value="")
xinchuan_var = tk.StringVar(value="")  # 是否信创，默认空
penetration_var = tk.StringVar(value="")  # 是否渗透，默认空
penetration_officer_var = tk.StringVar(value="")

tk.Label(root, text="选择输入文件夹：").grid(row=0, column=0, sticky="w", padx=10, pady=5)
tk.Entry(root, textvariable=folder_path, width=50).grid(row=0, column=1, padx=10, pady=5)
tk.Button(root, text="浏览...", command=select_folder).grid(row=0, column=2, padx=10, pady=5)

tk.Label(root, text="选择输出文件夹：").grid(row=1, column=0, sticky="w", padx=10, pady=5)
tk.Entry(root, textvariable=output_folder_path, width=50).grid(row=1, column=1, padx=10, pady=5)
tk.Button(root, text="浏览...", command=select_output_folder).grid(row=1, column=2, padx=10, pady=5)

# 是否渗透下拉框（倒数第三行）
tk.Label(root, text="是否渗透").grid(row=2, column=0, sticky="w", padx=10, pady=5)
penetration_cb = ttk.Combobox(root, textvariable=penetration_var, values=["", "渗透+扫描", "仅扫描", "放弃漏扫渗透"], state="readonly", width=30)
penetration_cb.grid(row=2, column=1, padx=10, pady=5)
penetration_var.trace_add("write", update_penetration_officer_row)

# 是否信创下拉框（倒数第二行）
tk.Label(root, text="是否信创").grid(row=3, column=0, sticky="w", padx=10, pady=5)
xinchuan_cb = ttk.Combobox(root, textvariable=xinchuan_var, values=["", "信创", "非信创"], state="readonly", width=30)
xinchuan_cb.grid(row=3, column=1, padx=10, pady=5)

# 地区下拉框（最后一行）
tk.Label(root, text="地区").grid(row=4, column=0, sticky="w", padx=10, pady=5)
region_cb = ttk.Combobox(root, textvariable=region_var, values=["", "浙江", "上海"], state="readonly", width=30)
region_cb.grid(row=4, column=1, padx=10, pady=5)

# 渗透师（初始隐藏）
penetration_officer_label = tk.Label(root, text="渗透师")
penetration_officer_cb = ttk.Combobox(root, textvariable=penetration_officer_var, values=penetration_officers, state="readonly", width=30)
penetration_officer_label.grid(row=5, column=0, sticky="w", padx=10, pady=5)
penetration_officer_cb.grid(row=5, column=1, padx=10, pady=5)
penetration_officer_label.grid_remove()
penetration_officer_cb.grid_remove()

entry_vars = []
check_vars = []

# 替换项输入框和勾选框，从第6行开始排布
start_row = 6
for i, item in enumerate(replace_items):
    row = start_row + i
    tk.Label(root, text=item["label"]).grid(row=row, column=0, sticky="w", padx=10, pady=5)

    if item["type"] == "text":
        var = tk.StringVar()
        entry = tk.Entry(root, textvariable=var, width=30)
        entry.grid(row=row, column=1, padx=10, pady=5)
        entry_vars.append(var)

    elif item["type"] == "select":
        var = tk.StringVar()
        cb = ttk.Combobox(root, textvariable=var, values=item["options"], state="readonly", width=28)
        cb.current(0)
        cb.grid(row=row, column=1, padx=10, pady=5)
        entry_vars.append(var)

    elif item["type"] == "date":
        var = tk.StringVar()
        entry = tk.Entry(root, textvariable=var, width=30)
        entry.grid(row=row, column=1, padx=10, pady=5)
        entry.bind("<Button-1>", lambda e, v=var, en=entry: open_calendar_popup(v, en))
        entry_vars.append(var)

    chk_var = tk.BooleanVar(value=True)
    tk.Checkbutton(root, text="启用", variable=chk_var).grid(row=row, column=3, padx=10, pady=5)
    check_vars.append(chk_var)

btn_row = start_row + len(replace_items)
tk.Button(root, text="开始替换", command=start_replace).grid(row=btn_row, column=1, pady=15)
tk.Button(root, text="全部勾选", command=lambda: [var.set(True) for var in check_vars]).grid(row=btn_row+1, column=1, pady=5)
tk.Button(root, text="取消全部勾选", command=lambda: [var.set(False) for var in check_vars]).grid(row=btn_row+2, column=1, pady=5)

root.mainloop()
